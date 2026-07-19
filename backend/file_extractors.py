"""
Multimodal file extractors for various document formats.
Unified interface for extracting text from different file types.
"""

import os
from typing import Optional, Tuple
from pathlib import Path


class FileExtractor:
    """
    Base class and factory for file extraction.
    Automatically detects file type and calls appropriate extractor.
    """

    EXTRACTORS = {}

    @classmethod
    def register(cls, extension: str):
        """Decorator to register a new extractor for a file type."""
        def wrapper(extractor_cls):
            cls.EXTRACTORS[extension.lower()] = extractor_cls
            return extractor_cls
        return wrapper

    @staticmethod
    def extract_text(file_path: str) -> Tuple[str, dict]:
        """
        Extract text from a file based on its extension.

        Args:
            file_path: Path to the file to extract from

        Returns:
            Tuple of (extracted_text, metadata)
            metadata includes: file_type, pages (if applicable), extraction_method, etc.

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
            Exception: If extraction fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = Path(file_path).suffix.lower()
        filename = os.path.basename(file_path)

        if file_ext not in FileExtractor.EXTRACTORS:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: {list(FileExtractor.EXTRACTORS.keys())}")

        extractor_cls = FileExtractor.EXTRACTORS[file_ext]
        extractor = extractor_cls()

        print(f"🔄 Extracting text from {filename} using {extractor_cls.__name__}...")
        text, metadata = extractor.extract(file_path)

        # Add common metadata
        metadata["filename"] = filename
        metadata["file_extension"] = file_ext
        metadata["file_size_bytes"] = os.path.getsize(file_path)

        print(f"✅ Successfully extracted {len(text)} characters from {filename}")

        return text, metadata


@FileExtractor.register(".txt")
class TextExtractor:
    """Extractor for plain text files."""

    def extract(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from a .txt file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()

        return text, {
            "file_type": "text",
            "extraction_method": "direct_read"
        }


@FileExtractor.register(".pdf")
class PDFExtractor:
    """Extractor for PDF files using pypdf."""

    def extract(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from a PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF extraction. Install it with: pip install pypdf")

        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                # Add page marker for reference
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        return full_text, {
            "file_type": "pdf",
            "total_pages": total_pages,
            "extraction_method": "pypdf"
        }


@FileExtractor.register(".docx")
class DocxExtractor:
    """Extractor for Word (.docx) files."""

    def extract(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from a .docx file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for Word extraction. Install it with: pip install python-docx")

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(table_text)

        full_text = "\n\n".join(text_parts)

        return full_text, {
            "file_type": "docx",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "extraction_method": "python-docx"
        }

    @staticmethod
    def _extract_table(table) -> str:
        """Convert a Word table to text format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


@FileExtractor.register(".xlsx")
class ExcelExtractor:
    """Extractor for Excel (.xlsx) files using openpyxl."""

    def extract(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from an Excel file."""
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for Excel extraction. Install it with: pip install openpyxl")

        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_parts = []
        sheet_count = 0

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_count += 1

            # Add sheet header
            text_parts.append(f"=== Sheet: {sheet_name} ===")

            # Extract all rows
            for row in sheet.iter_rows(values_only=True):
                # Filter out None values and convert to string
                row_values = [str(cell) if cell is not None else "" for cell in row]
                # Skip empty rows
                if any(row_values):
                    row_text = " | ".join(row_values)
                    text_parts.append(row_text)

            text_parts.append("")  # Add spacing between sheets

        full_text = "\n".join(text_parts)

        return full_text, {
            "file_type": "xlsx",
            "sheets": sheet_count,
            "sheet_names": workbook.sheetnames,
            "extraction_method": "openpyxl"
        }


# Alias for backward compatibility
class MultimodalExtractor:
    """Alias for FileExtractor."""
    @staticmethod
    def extract_text(file_path: str) -> Tuple[str, dict]:
        return FileExtractor.extract_text(file_path)


# Test function
if __name__ == "__main__":
    """Quick test of extractors."""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python file_extractors.py <file_path>")
        print("\nSupported formats:", list(FileExtractor.EXTRACTORS.keys()))
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        text, metadata = FileExtractor.extract_text(file_path)
        print("\n" + "=" * 60)
        print(f"Metadata: {metadata}")
        print("=" * 60)
        print(f"Extracted text ({len(text)} chars):\n")
        print(text[:500] + "..." if len(text) > 500 else text)
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
