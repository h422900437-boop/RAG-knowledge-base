#!/usr/bin/env python3
"""
Test script for multimodal file extraction.
Tests PDF, Word, Excel extraction capabilities.
"""

import os
import sys
from pathlib import Path

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from backend.file_extractors import FileExtractor


def test_extractor(file_path: str):
    """Test extracting text from a file."""
    print(f"\n{'=' * 70}")
    print(f"Testing: {os.path.basename(file_path)}")
    print(f"{'=' * 70}\n")

    try:
        text, metadata = FileExtractor.extract_text(file_path)

        print("✅ Extraction successful!")
        print(f"\n📊 Metadata:")
        for key, value in metadata.items():
            print(f"  - {key}: {value}")

        print(f"\n📝 Extracted text (first 500 characters):\n")
        print(text[:500])
        if len(text) > 500:
            print("\n... (truncated)")

        print(f"\n📈 Total extracted: {len(text)} characters")
        return True

    except FileNotFoundError as e:
        print(f"❌ File not found: {e}")
        return False
    except Exception as e:
        print(f"❌ Extraction failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def create_sample_files():
    """Create sample files for testing in the data directory."""
    data_dir = Path(__file__).parent.parent / "data"
    data_dir.mkdir(exist_ok=True)

    # Sample PDF creation (requires pypdf or similar)
    # Note: This is a simplified example. For real testing, use actual files.

    # Sample Word document creation
    try:
        from docx import Document

        doc_path = data_dir / "sample_policy.docx"
        if not doc_path.exists():
            doc = Document()
            doc.add_heading('Company Policy Document', 0)
            doc.add_heading('Section 1: Attendance Policy', level=1)
            doc.add_paragraph('Employees are expected to arrive on time for their scheduled shifts.')
            doc.add_paragraph('Absence without notice will be recorded as unexcused absence.')

            doc.add_heading('Section 2: Leave Policy', level=1)
            doc.add_paragraph('Employees are entitled to annual leave as per company policy.')

            # Add a sample table
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Leave Type'
            hdr_cells[1].text = 'Duration'
            hdr_cells[2].text = 'Notes'

            row_cells = table.rows[1].cells
            row_cells[0].text = 'Annual Leave'
            row_cells[1].text = '20 days'
            row_cells[2].text = 'Paid leave'

            row_cells = table.rows[2].cells
            row_cells[0].text = 'Sick Leave'
            row_cells[1].text = '10 days'
            row_cells[2].text = 'Medical certificate required'

            doc.save(doc_path)
            print(f"✅ Created sample Word file: {doc_path}")

    except ImportError:
        print("⚠️  python-docx not installed. Skipping Word document creation.")

    # Sample Excel document creation
    try:
        import openpyxl

        xlsx_path = data_dir / "sample_data.xlsx"
        if not xlsx_path.exists():
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Employee Data"

            # Headers
            ws['A1'] = 'Employee ID'
            ws['B1'] = 'Name'
            ws['C1'] = 'Department'
            ws['D1'] = 'Salary'

            # Sample data
            ws['A2'] = '001'
            ws['B2'] = 'John Doe'
            ws['C2'] = 'HR'
            ws['D2'] = '50000'

            ws['A3'] = '002'
            ws['B3'] = 'Jane Smith'
            ws['C3'] = 'IT'
            ws['D3'] = '60000'

            ws['A4'] = '003'
            ws['B4'] = 'Bob Johnson'
            ws['C4'] = 'Sales'
            ws['D4'] = '45000'

            wb.save(xlsx_path)
            print(f"✅ Created sample Excel file: {xlsx_path}")

    except ImportError:
        print("⚠️  openpyxl not installed. Skipping Excel document creation.")

    return data_dir


def main():
    """Run all tests."""
    print("🧪 Multimodal File Extraction Test Suite\n")

    # Check supported formats
    print("📋 Supported formats:")
    for ext in FileExtractor.EXTRACTORS.keys():
        print(f"  - {ext}")
    print()

    # Create sample files
    data_dir = create_sample_files()

    # Test existing file
    existing_file = data_dir / "company_policies.txt"
    if existing_file.exists():
        test_extractor(str(existing_file))
    else:
        print(f"ℹ️  {existing_file} not found, skipping TXT test.")

    # Test Word file
    docx_file = data_dir / "sample_policy.docx"
    if docx_file.exists():
        test_extractor(str(docx_file))
    else:
        print(f"\nℹ️  {docx_file} not found, create it first.")

    # Test Excel file
    xlsx_file = data_dir / "sample_data.xlsx"
    if xlsx_file.exists():
        test_extractor(str(xlsx_file))
    else:
        print(f"\nℹ️  {xlsx_file} not found, create it first.")

    # Test PDF file if available
    pdf_files = list(data_dir.glob("*.pdf"))
    if pdf_files:
        for pdf_file in pdf_files[:1]:  # Test first PDF only
            test_extractor(str(pdf_file))
    else:
        print(f"\nℹ️  No PDF files found in {data_dir}")

    print("\n" + "=" * 70)
    print("✅ Test suite complete!")
    print("=" * 70)


if __name__ == "__main__":
    main()
